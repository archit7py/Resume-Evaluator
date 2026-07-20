import os
from pydantic import BaseModel
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("GROQ_API_KEY environment variable is not set.")

client = Groq(api_key=my_api_key)

model = "llama-3.3-70b-versatile"

job_description = """
Project Role : AI / ML Engineer
Project Role Description : Develops applications and systems that utilize AI tools, Cloud AI services, with proper cloud or on-prem application pipeline with production ready quality. Be able to apply GenAI models as part of the solution. Could also include but not limited to deep learning, neural networks, chatbots, image processing.
Must have skills : Machine Learning (ML)
Good to have skills : NA
Minimum 2 year(s) of experience is required
Educational Qualification : 15 years full time education

Summary:
These roles have many overlapping skills with GENAI Engineers and architects. Description may scaleup/scale down based on expected seniority.

Roles & Responsibilities:
-Implement generative AI models, identify insights that can be used to drive business decisions. Work closely with multi-functional teams to understand business problems, develop hypotheses, and test those hypotheses with data, collaborating with cross-functional teams to define AI project requirements and objectives, ensuring alignment with overall business goals.
-Conducting research to stay up-to-date with the latest advancements in generative AI, machine learning, and deep learning techniques and identify opportunities to integrate them into our products and services.
-Optimizing existing generative AI models for improved performance, scalability, and efficiency.
-Ensure data quality and accuracy
-Leading the design and development of prompt engineering strategies and techniques to optimize the performance and output of our GenAI models.
-Implementing cutting-edge NLP techniques and prompt engineering methodologies to enhance the capabilities and efficiency of our GenAI models.
-Determining the most effective prompt generation processes and approaches to drive innovation and excellence in the field of AI technology, collaborating with AI researchers and developers
-Experience working with cloud based platforms (example: AWS, Azure or related)
-Strong problem-solving and analytical skills
-Proficiency in handling various data formats and sources through Omni Channel for Speech and voice applications, part of conversational AI
-Prior statistical modelling experience
-Demonstrable experience with deep learning algorithms and neural networks
-Developing clear and concise documentation, including technical specifications, user guides, and presentations, to communicate complex AI concepts to both technical and non-technical stakeholders.
-Contributing to the establishment of best practices and standards for generative AI development within the organization.

Professional & Technical Skills:
-Must have solid experience developing and implementing generative AI models, with a strong understanding of deep learning techniques such as GPT, VAE, and GANs.
-Must be proficient in Python and have experience with machine learning libraries and frameworks such as TensorFlow, PyTorch, or Keras.
-Must have strong knowledge of data structures, algorithms, and software engineering principles.
-Must be familiar with cloud-based platforms and services, such as AWS, GCP, or Azure.
-Need to have experience with natural language processing (NLP) techniques and tools, such as SpaCy, NLTK, or Hugging Face.
-Must be familiar with data visualization tools and libraries, such as Matplotlib, Seaborn, or Plotly.
-Need to have knowledge of software development methodologies, such as Agile or Scrum.
-Possess excellent problem-solving skills, with the ability to think critically and creatively to develop innovative AI solutions.

Additional Information:
-Must have a degree in Computer Science, Artificial Intelligence, Machine Learning, or a related field. A Ph.D. is highly desirable.
-strong communication skills, with the ability to effectively convey complex technical concepts to a diverse audience.
-You possess a proactive mindset, with the ability to work independently and collaboratively in a fast-paced, dynamic environment.

Additional Information:
- The candidate should have a minimum of 5 years of real time experience in Machine Learning.
- A 15 years full time education is required.

"""


# ---------------------------------------------------------------------------
# Job description schema
# ---------------------------------------------------------------------------
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requiremnets: list[str]
    responsiblities: list[str]


jobd_schema = JobD.model_json_schema()

jd_system_prompt = f"""
You are an expert HR assistant.
Your job is to analyze the job description and extract the required information based on the schema and return it in the json output format.{jobd_schema}

IMPORTANT
1. Return ONLY valid JSON.
2. Do NOT return Markdown.
3. Do NOT wrap the JSON inside ```json ... ```.
4. Do NOT explain your answer.
5. Do NOT add extra fields that are not present in the schema.
6. Follow the schema EXACTLY.
7. Every key in the schema must be present in the output.
8. If a value is not found in the resume, return:
   - "" for strings
   - [] for lists
   - 0 for numeric values
9. Never guess, infer, or hallucinate information.
10. Extract only information explicitly mentioned in the resume.
11. Preserve original spellings of names, companies, universities, certifications, and technologies.
12. Remove duplicate entries from lists.
13. Skills must be returned as a list of unique strings.
14. Projects must be returned as a list.
15. Certifications must be returned as a list.
16. Experience should be extracted exactly as mentioned.
17. Education details should be extracted exactly as written.
18. Ignore decorative text, page numbers, headers, footers, icons, and formatting.
19. Ignore irrelevant information that is not part of the candidate's professional profile.
20. Do not create values for missing information.

You MUST validate your response against the following JSON schema before returning it.
"""

jd_user_prompt = f"""
Analyze the following job description job_description.
{job_description}
"""


def call_groq(system_prompt: str, user_prompt: str) -> dict:
    """Helper that sends a system/user prompt pair to Groq and returns parsed JSON."""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0,
        response_format={"type": "json_object"},
    )
    raw_json = response.choices[0].message.content
    return json.loads(raw_json)


import json

job_data = call_groq(jd_system_prompt, jd_user_prompt)
job = JobD(**job_data)
print("=== Job Description Parsed ===")
print(job.role)
print(job.minimum_experience)
print(job.education_requiremnets)


# ---------------------------------------------------------------------------
# Resume schema
# ---------------------------------------------------------------------------
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()

resume_system_prompt = f"""
You are an expert HR assistant.
Your job is to analyze the candidate's resume and extract the required information based on the schema and return it in the json output format.{resume_schema}

IMPORTANT
1. Return ONLY valid JSON.
2. Do NOT return Markdown.
3. Do NOT wrap the JSON inside ```json ... ```.
4. Do NOT explain your answer.
5. Do NOT add extra fields that are not present in the schema.
6. Follow the schema EXACTLY.
7. Every key in the schema must be present in the output.
8. If a value is not found in the resume, return:
   - "" for strings
   - [] for lists
   - 0 for numeric values
   - null for optional fields with no value
9. Never guess, infer, or hallucinate information.
10. Extract only information explicitly mentioned in the resume.
11. Preserve original spellings of names, companies, universities, certifications, and technologies.
12. Remove duplicate entries from lists.
13. Skills must be returned as a list of unique strings.
14. Projects must be returned as a list.
15. Certifications must be returned as a list.
16. Experience should be extracted exactly as mentioned.
17. Education details should be extracted exactly as written.
18. Ignore decorative text, page numbers, headers, footers, icons, and formatting.
19. Ignore irrelevant information that is not part of the candidate's professional profile.
20. Do not create values for missing information.

You MUST validate your response against the following JSON schema before returning it.
"""


def parse_resume(resume_text: str) -> Resume:
    """Send raw resume text to the LLM and parse the structured result."""
    user_prompt = f"""
    Analyze the following resume text and extract the candidate's information.
    {resume_text}
    """
    data = call_groq(resume_system_prompt, user_prompt)
    return Resume(**data)


def _extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages_text).strip()

    if not text:
        raise ValueError(
            f"No extractable text found in '{path.name}'. "
            "This usually means the PDF is a scanned image rather than "
            "a text-based PDF, so it needs OCR before it can be parsed."
        )
    return text


def _extract_docx_text(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Resumes often put skills/experience in tables — grab those too.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(f"No extractable text found in '{path.name}'.")
    return text


def load_resume_text(path: str) -> str:
    """
    Load raw text from a resume file.
    Auto-detects the file type from its extension:
      - .pdf  -> extracted with pypdf
      - .docx -> extracted with python-docx (paragraphs + tables)
      - .txt  -> read directly
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(p)
    elif suffix == ".docx":
        return _extract_docx_text(p)
    elif suffix == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported resume file type: '{suffix}'. "
            "Supported types are .pdf, .docx, and .txt."
        )



# ---------------------------------------------------------------------------
# Matching logic
# ---------------------------------------------------------------------------
class MatchResult(BaseModel):
    score: float
    details: dict


def match_resume_to_job(resume: Resume, job: JobD) -> MatchResult:
    """
    Very simple, transparent scoring model:
    - 60% weight on required-skill overlap
    - 20% weight on preferred-skill overlap
    - 20% weight on meeting minimum experience
    Swap this out for embeddings / a fuzzy matcher if you want something smarter.
    """
    resume_skills = {s.strip().lower() for s in resume.skills if s}

    required = {s.strip().lower() for s in job.required_skills if s}
    preferred = {s.strip().lower() for s in job.preferred_skills if s}

    matched_required = resume_skills & required
    matched_preferred = resume_skills & preferred

    required_score = (len(matched_required) / len(required)) if required else 1.0
    preferred_score = (len(matched_preferred) / len(preferred)) if preferred else 1.0

    experience_score = 1.0
    if job.minimum_experience:
        candidate_years = resume.total_experience_years or 0.0
        experience_score = min(candidate_years / job.minimum_experience, 1.0)

    final_score = round(
        (required_score * 0.6 + preferred_score * 0.2 + experience_score * 0.2) * 100,
        2,
    )

    details = {
        "matched_required_skills": sorted(matched_required),
        "missing_required_skills": sorted(required - resume_skills),
        "matched_preferred_skills": sorted(matched_preferred),
        "missing_preferred_skills": sorted(preferred - resume_skills),
        "required_score_pct": round(required_score * 100, 2),
        "preferred_score_pct": round(preferred_score * 100, 2),
        "experience_score_pct": round(experience_score * 100, 2),
        "candidate_experience_years": resume.total_experience_years,
        "required_experience_years": job.minimum_experience,
    }

    return MatchResult(score=final_score, details=details)


def get_resume_files_from_folder(folder_path) -> list[str]:
    """
    Scan a folder and return paths to all supported resume files
    (.pdf, .docx, .txt), sorted alphabetically. Ignores everything else
    (e.g. .py, .toml, .md project files) so it's safe to point at a
    project folder, not just a dedicated resumes-only folder.
    """
    folder = Path(folder_path)
    if not folder.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")
    if not folder.is_dir():
        raise NotADirectoryError(f"Not a folder: {folder}")

    supported_extensions = {".pdf", ".docx", ".txt"}
    files = sorted(
        str(f) for f in folder.iterdir()
        if f.is_file() and f.suffix.lower() in supported_extensions
    )
    return files


# ---------------------------------------------------------------------------
# Example run
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import sys

    # Usage:
    #   python resume_parser.py                     -> looks in the "resumes" folder next to this script
    #   python resume_parser.py path/to/folder       -> looks in that folder instead
    #   python resume_parser.py resume1.pdf r2.docx  -> specific files (old behavior still works)
    #
    # NOTE: RESUMES_FOLDER is resolved relative to THIS SCRIPT'S location,
    # not your current working directory. So it works the same whether you run
    # `python resume_parser.py` from this folder or `uv run resume_parser.py`
    # from somewhere else.
    RESUMES_FOLDER = Path(__file__).resolve().parent / "resumes"

    args = sys.argv[1:]

    if len(args) == 1 and Path(args[0]).is_dir():
        # A single folder path was passed.
        resume_paths = get_resume_files_from_folder(args[0])
    elif args:
        # One or more individual file paths were passed.
        resume_paths = args
    else:
        # No arguments -> default to the "resumes" folder next to this script.
        try:
            resume_paths = get_resume_files_from_folder(RESUMES_FOLDER)
            if not resume_paths:
                print(
                    f"'{RESUMES_FOLDER}' exists but has no .pdf/.docx/.txt files in it. "
                    "Add resumes there and run again."
                )
        except FileNotFoundError:
            print(
                f"No '{RESUMES_FOLDER}' folder found next to this script. "
                "Create a 'resumes' folder here and put your resume files in it, "
                "or run: python resume_parser.py path/to/your/folder"
            )
            resume_paths = []

    results = []  # list of (path, Resume, MatchResult)

    if resume_paths:
        print(f"Found {len(resume_paths)} resume(s) to process.")
        for resume_path in resume_paths:
            print(f"\n=== Processing: {resume_path} ===")
            try:
                resume_text = load_resume_text(resume_path)
                resume = parse_resume(resume_text)
                match_result = match_resume_to_job(resume, job)
                results.append((resume_path, resume, match_result))

                print(f"Name: {resume.name}")
                print(f"Experience: {resume.total_experience_years} years")
                print(f"Skills: {resume.skills}")
                print(f"Score: {match_result.score}%")
            except (FileNotFoundError, ValueError) as e:
                print(f"Skipped '{resume_path}': {e}")

        # Rank all successfully-processed resumes by match score, best first.
        results.sort(key=lambda r: r[2].score, reverse=True)

        print("\n=== Ranking (best match first) ===")
        for rank, (path, resume, match_result) in enumerate(results, start=1):
            print(f"{rank}. {resume.name or path} - {match_result.score}%")

    else:
        # Fallback sample so the script still runs with nothing configured.
        resume_path = "(sample text - no file passed)"
        resume_text = """
        John Doe
        john.doe@email.com | +1-555-123-4567

        Summary:
        Software engineer with 4 years of experience building backend services in Python and Java.

        Experience:
        Software Development Engineer, Acme Corp, Jan 2021 - Present
        - Designed and built scalable microservices handling high-throughput traffic.
        - Worked on system architecture and reliability improvements.
        Skills used: Python, AWS, Docker, System Design

        Software Engineer, Beta Inc, Jun 2019 - Dec 2020
        - Built internal tooling in Java and Python.
        Skills used: Java, Python, SQL

        Education:
        B.Tech in Computer Science, XYZ University, 2019

        Skills: Python, Java, AWS, Docker, System Design, SQL, Git

        Certifications:
        AWS Certified Developer - Associate
        """

        print(f"\n=== Extracting resume from: {resume_path} ===")
        resume = parse_resume(resume_text)
        print("\n=== Resume Parsed ===")
        print(resume.name)
        print(resume.total_experience_years)
        print(resume.skills)

        result = match_resume_to_job(resume, job)
        print("\n=== Match Result ===")
        print(f"Score: {result.score}%")
        print(result.details)